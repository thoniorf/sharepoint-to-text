"""
DOCX content extractor using python-docx library.
"""

import datetime
import io
import logging
import zipfile
from typing import Any, Generator
from xml.etree import ElementTree as ET

from docx import Document
from docx.document import Document as DocumentObject
from docx.oxml.ns import qn

from sharepoint2text.extractors.data_types import (
    DocxComment,
    DocxContent,
    DocxFormula,
    DocxHeaderFooter,
    DocxHyperlink,
    DocxImage,
    DocxMetadata,
    DocxNote,
    DocxParagraph,
    DocxRun,
    DocxSection,
)

logger = logging.getLogger(__name__)


class _DocxFullTextExtractor:
    """Extracts a full text representation of a docx file.
    Respects paragraphs, tables and formulas and their order of occurrence."""

    # Greek letter and symbol mapping for LaTeX conversion
    GREEK_TO_LATEX = {
        # Lowercase Greek
        "α": "\\alpha",
        "β": "\\beta",
        "γ": "\\gamma",
        "δ": "\\delta",
        "ε": "\\epsilon",
        "ζ": "\\zeta",
        "η": "\\eta",
        "θ": "\\theta",
        "ι": "\\iota",
        "κ": "\\kappa",
        "λ": "\\lambda",
        "μ": "\\mu",
        "ν": "\\nu",
        "ξ": "\\xi",
        "ο": "o",  # omicron is just 'o' in LaTeX
        "π": "\\pi",
        "ρ": "\\rho",
        "σ": "\\sigma",
        "ς": "\\varsigma",
        "τ": "\\tau",
        "υ": "\\upsilon",
        "φ": "\\phi",
        "χ": "\\chi",
        "ψ": "\\psi",
        "ω": "\\omega",
        # Uppercase Greek
        "Α": "A",
        "Β": "B",
        "Γ": "\\Gamma",
        "Δ": "\\Delta",
        "Ε": "E",
        "Ζ": "Z",
        "Η": "H",
        "Θ": "\\Theta",
        "Ι": "I",
        "Κ": "K",
        "Λ": "\\Lambda",
        "Μ": "M",
        "Ν": "N",
        "Ξ": "\\Xi",
        "Ο": "O",
        "Π": "\\Pi",
        "Ρ": "P",
        "Σ": "\\Sigma",
        "Τ": "T",
        "Υ": "\\Upsilon",
        "Φ": "\\Phi",
        "Χ": "X",
        "Ψ": "\\Psi",
        "Ω": "\\Omega",
        # Common math symbols
        "∞": "\\infty",
        "∂": "\\partial",
        "∇": "\\nabla",
        "±": "\\pm",
        "∓": "\\mp",
        "×": "\\times",
        "÷": "\\div",
        "·": "\\cdot",
        "≤": "\\leq",
        "≥": "\\geq",
        "≠": "\\neq",
        "≈": "\\approx",
        "≡": "\\equiv",
        "∈": "\\in",
        "∉": "\\notin",
        "⊂": "\\subset",
        "⊃": "\\supset",
        "⊆": "\\subseteq",
        "⊇": "\\supseteq",
        "∪": "\\cup",
        "∩": "\\cap",
        "∧": "\\land",
        "∨": "\\lor",
        "¬": "\\neg",
        "→": "\\rightarrow",
        "←": "\\leftarrow",
        "↔": "\\leftrightarrow",
        "⇒": "\\Rightarrow",
        "⇐": "\\Leftarrow",
        "⇔": "\\Leftrightarrow",
        "∀": "\\forall",
        "∃": "\\exists",
        "∅": "\\emptyset",
        "ℕ": "\\mathbb{N}",
        "ℤ": "\\mathbb{Z}",
        "ℚ": "\\mathbb{Q}",
        "ℝ": "\\mathbb{R}",
        "ℂ": "\\mathbb{C}",
    }

    @classmethod
    def _convert_greek_and_symbols(cls, text: str) -> str:
        """Convert Greek letters and math symbols to LaTeX equivalents."""
        result = []
        for char in text:
            if char in cls.GREEK_TO_LATEX:
                result.append(cls.GREEK_TO_LATEX[char])
            else:
                result.append(char)
        return "".join(result)

    @classmethod
    def omml_to_latex(cls, omath_element) -> str:
        """Convert OMML element to LaTeX-like string.

        Handles malformed bracket placement in sqrt/rad elements by consuming
        content until the matching closing bracket is found.
        """
        m_ns = "{http://schemas.openxmlformats.org/officeDocument/2006/math}"
        parts = []
        pending_sqrt_close = None  # Bracket needed to close current sqrt

        # Property elements to skip
        skip_tags = {
            "rPr",
            "fPr",
            "radPr",
            "ctrlPr",
            "oMathParaPr",
            "degHide",
            "type",
            "rFonts",
            "i",
            "color",
            "sz",
            "szCs",
            "jc",
            "solidFill",
            "srgbClr",
            "latin",
        }

        def process_element(elem) -> str:
            """Recursively process an element and return its LaTeX representation."""
            nonlocal pending_sqrt_close

            if elem is None:
                return ""

            tag = elem.tag.split("}")[-1]

            # Skip property elements
            if tag in skip_tags:
                return ""

            # Text content (both w:t and m:t)
            if tag == "t":
                text = elem.text or ""
                converted = cls._convert_greek_and_symbols(text)

                # Handle malformed sqrt: if we're waiting for a closing bracket
                if pending_sqrt_close and pending_sqrt_close in converted:
                    idx = converted.index(pending_sqrt_close)
                    inside = converted[:idx]  # Content inside sqrt
                    outside = converted[idx + 1 :]  # Content after closing bracket
                    pending_sqrt_close = None
                    return inside + "}" + outside

                return converted

            # Fraction: m:f contains m:num (numerator) and m:den (denominator)
            if tag == "f":
                num = elem.find(f"{m_ns}num")
                den = elem.find(f"{m_ns}den")
                num_text = process_element(num)
                den_text = process_element(den)
                return f"\\frac{{{num_text}}}{{{den_text}}}"

            # Superscript: m:sSup contains m:e (base) and m:sup (superscript)
            if tag == "sSup":
                base = elem.find(f"{m_ns}e")
                sup = elem.find(f"{m_ns}sup")
                base_text = process_element(base)
                sup_text = process_element(sup)
                return f"{base_text}^{{{sup_text}}}"

            # Subscript: m:sSub contains m:e (base) and m:sub (subscript)
            if tag == "sSub":
                base = elem.find(f"{m_ns}e")
                sub = elem.find(f"{m_ns}sub")
                base_text = process_element(base)
                sub_text = process_element(sub)
                return f"{base_text}_{{{sub_text}}}"

            # Sub-superscript: m:sSubSup contains m:e, m:sub, and m:sup
            if tag == "sSubSup":
                base = elem.find(f"{m_ns}e")
                sub = elem.find(f"{m_ns}sub")
                sup = elem.find(f"{m_ns}sup")
                base_text = process_element(base)
                sub_text = process_element(sub)
                sup_text = process_element(sup)
                return f"{base_text}_{{{sub_text}}}^{{{sup_text}}}"

            # Square root: m:rad contains m:deg (degree, optional) and m:e (content)
            if tag == "rad":
                deg = elem.find(f"{m_ns}deg")
                content = elem.find(f"{m_ns}e")
                content_text = process_element(content)
                deg_text = process_element(deg).strip()

                # Handle malformed: lone opening bracket inside sqrt
                # Some OMML has sqrt containing just "(" with content after
                if content_text.strip() in ("(", "[", "{"):
                    bracket_map = {"(": ")", "[": "]", "{": "}"}
                    pending_sqrt_close = bracket_map.get(content_text.strip(), ")")
                    if deg_text:
                        return f"\\sqrt[{deg_text}]{{"
                    else:
                        return "\\sqrt{"
                else:
                    if deg_text:
                        return f"\\sqrt[{deg_text}]{{{content_text}}}"
                    else:
                        return f"\\sqrt{{{content_text}}}"

            # N-ary (sum, product, integral): m:nary
            if tag == "nary":
                chr_elem = elem.find(f".//{m_ns}chr")
                op = chr_elem.get(f"{m_ns}val") if chr_elem is not None else "∑"

                sub = elem.find(f"{m_ns}sub")
                sup = elem.find(f"{m_ns}sup")
                content = elem.find(f"{m_ns}e")

                op_map = {
                    "∑": "\\sum",
                    "∏": "\\prod",
                    "∫": "\\int",
                    "∬": "\\iint",
                    "∭": "\\iiint",
                }
                latex_op = op_map.get(op, cls._convert_greek_and_symbols(op))

                sub_text = process_element(sub)
                sup_text = process_element(sup)
                content_text = process_element(content)

                result = latex_op
                if sub_text.strip():
                    result += f"_{{{sub_text}}}"
                if sup_text.strip():
                    result += f"^{{{sup_text}}}"
                result += f" {content_text}"
                return result

            # Delimiter (parentheses, brackets): m:d
            if tag == "d":
                beg_chr = elem.find(f".//{m_ns}begChr")
                end_chr = elem.find(f".//{m_ns}endChr")
                left = beg_chr.get(f"{m_ns}val") if beg_chr is not None else "("
                right = end_chr.get(f"{m_ns}val") if end_chr is not None else ")"

                e_elements = elem.findall(f"{m_ns}e")
                content_parts = [process_element(e) for e in e_elements]
                content_text = ", ".join(content_parts)
                return f"{left}{content_text}{right}"

            # Matrix: m:m contains m:mr (rows) which contain m:e (elements)
            if tag == "m" and elem.find(f"{m_ns}mr") is not None:
                rows = []
                for mr in elem.findall(f"{m_ns}mr"):
                    cells = [process_element(e) for e in mr.findall(f"{m_ns}e")]
                    rows.append(" & ".join(cells))
                return "\\begin{matrix}" + " \\\\ ".join(rows) + "\\end{matrix}"

            # Function: m:func contains m:fName and m:e
            if tag == "func":
                fname = elem.find(f"{m_ns}fName")
                content = elem.find(f"{m_ns}e")
                fname_text = process_element(fname)
                content_text = process_element(content)
                func_map = {
                    "sin": "\\sin",
                    "cos": "\\cos",
                    "tan": "\\tan",
                    "log": "\\log",
                    "ln": "\\ln",
                    "lim": "\\lim",
                    "exp": "\\exp",
                    "max": "\\max",
                    "min": "\\min",
                }
                latex_fname = func_map.get(fname_text.strip(), fname_text)
                return f"{latex_fname}{{{content_text}}}"

            # Bar/overline: m:bar
            if tag == "bar":
                content = elem.find(f"{m_ns}e")
                content_text = process_element(content)
                return f"\\overline{{{content_text}}}"

            # Accent (hat, tilde, etc.): m:acc
            if tag == "acc":
                chr_elem = elem.find(f".//{m_ns}chr")
                accent = chr_elem.get(f"{m_ns}val") if chr_elem is not None else "^"
                content = elem.find(f"{m_ns}e")
                content_text = process_element(content)

                accent_map = {
                    "̂": "\\hat",
                    "̃": "\\tilde",
                    "̄": "\\bar",
                    "⃗": "\\vec",
                    "̇": "\\dot",
                }
                latex_accent = accent_map.get(accent, "\\hat")
                return f"{latex_accent}{{{content_text}}}"

            # Default: recurse into children and concatenate results
            result = []
            for child in elem:
                child_result = process_element(child)
                if child_result:
                    result.append(child_result)
            return "".join(result)

        # Process all children of the omath element
        for child in omath_element:
            child_result = process_element(child)
            if child_result:
                parts.append(child_result)

        # If sqrt was never closed (no matching bracket found), close it now
        if pending_sqrt_close:
            parts.append("}")

        return "".join(parts)

    @classmethod
    def extract_full_text(
        cls, file_like: io.BytesIO, include_formulas: bool = True
    ) -> str:
        """Combines the full text of the docx file into a single text.
        Paragraphs, tables, and equations are kept in the order of occurrence.

        Args:
            file_like: BytesIO object containing the DOCX file
            include_formulas: Whether to include LaTeX formulas in output (default: True)
        """
        logger.debug("Extracting document full text")
        file_like.seek(0)
        doc = Document(file_like)
        all_text = []

        w_ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
        m_ns = "{http://schemas.openxmlformats.org/officeDocument/2006/math}"
        mc_ns = "{http://schemas.openxmlformats.org/markup-compatibility/2006}"

        def process_element(elem, parts: list):
            """Recursively process element, handling AlternateContent properly.

            Only processes mc:Choice content and skips mc:Fallback to avoid
            extracting duplicate content from fallback representations.
            """
            tag = elem.tag.split("}")[-1]

            # Handle AlternateContent - only use Choice, skip Fallback
            if tag == "AlternateContent":
                choice = elem.find(f"{mc_ns}Choice")
                if choice is not None:
                    for child in choice:
                        process_element(child, parts)
                return

            # Skip Fallback elements entirely to avoid duplicate content
            if tag == "Fallback":
                return

            # Regular run of text
            if tag == "r":
                for child in elem:
                    child_tag = child.tag.split("}")[-1]
                    if child_tag == "t":
                        if child.text:
                            parts.append(child.text)
                    elif child_tag == "AlternateContent":
                        process_element(child, parts)
                return

            # Inline equation
            if tag == "oMath":
                if include_formulas:
                    latex = cls.omml_to_latex(elem)
                    if latex.strip():
                        parts.append(f"${latex}$")
                return

            # Display equation
            if tag == "oMathPara":
                if include_formulas:
                    omath = elem.find(f"{m_ns}oMath")
                    if omath is not None:
                        latex = cls.omml_to_latex(omath)
                        if latex.strip():
                            parts.append(f"$${latex}$$")
                return

            # Recurse into other elements
            for child in elem:
                process_element(child, parts)

        def extract_paragraph_content(p_element) -> str:
            """Extract text from paragraph including inline and display equations."""
            parts = []
            for child in p_element:
                process_element(child, parts)
            return "".join(parts)

        def extract_table_text(tbl_element) -> list[str]:
            """Extract text from table in row order."""
            texts = []
            for row in tbl_element.iter(f"{w_ns}tr"):
                for cell in row.iter(f"{w_ns}tc"):
                    cell_parts = []
                    for p in cell.iter(f"{w_ns}p"):
                        text = extract_paragraph_content(p)
                        if text.strip():
                            cell_parts.append(text)
                    if cell_parts:
                        texts.append(" ".join(cell_parts))
            return texts

        # Iterate through body elements in document order
        for element in doc.element.body:
            tag = element.tag.split("}")[-1]

            if tag == "p":  # Paragraph (may contain oMathPara)
                text = extract_paragraph_content(element)
                if text.strip():
                    all_text.append(text)

            elif tag == "tbl":  # Table
                table_texts = extract_table_text(element)
                all_text.extend(table_texts)

        return "\n".join(all_text)


def _extract_footnotes(file_like: io.BytesIO) -> list[DocxNote]:
    logger.debug("Extracting footnotes")
    footnotes = []
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    w_ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

    with zipfile.ZipFile(file_like, "r") as z:
        if "word/footnotes.xml" not in z.namelist():
            return footnotes

        with z.open("word/footnotes.xml") as f:
            tree = ET.parse(f)

        for fn in tree.findall(".//w:footnote", ns):
            fn_id = fn.get(f"{w_ns}id") or ""
            if fn_id not in ["-1", "0"]:  # Skip separator and continuation footnotes
                footnotes.append(
                    DocxNote(
                        id=fn_id,
                        text="".join(t.text or "" for t in fn.findall(".//w:t", ns)),
                    )
                )

    return footnotes


def _extract_comments(file_like: io.BytesIO) -> list[DocxComment]:
    logger.debug("Extracting comments")
    file_like.seek(0)
    comments = []
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    w_ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

    with zipfile.ZipFile(file_like, "r") as z:
        if "word/comments.xml" not in z.namelist():
            return comments

        with z.open("word/comments.xml") as f:
            tree = ET.parse(f)

        for comment in tree.findall(".//w:comment", ns):
            comments.append(
                DocxComment(
                    id=comment.get(f"{w_ns}id") or "",
                    author=comment.get(f"{w_ns}author") or "",
                    date=comment.get(f"{w_ns}date") or "",
                    text="".join(t.text or "" for t in comment.findall(".//w:t", ns)),
                )
            )

    return comments


def _extract_sections(doc: DocumentObject) -> list[DocxSection]:
    logger.debug("Extracting sections")
    sections = []
    for section in doc.sections:
        sections.append(
            DocxSection(
                page_width_inches=(
                    section.page_width.inches if section.page_width else None
                ),
                page_height_inches=(
                    section.page_height.inches if section.page_height else None
                ),
                left_margin_inches=(
                    section.left_margin.inches if section.left_margin else None
                ),
                right_margin_inches=(
                    section.right_margin.inches if section.right_margin else None
                ),
                top_margin_inches=(
                    section.top_margin.inches if section.top_margin else None
                ),
                bottom_margin_inches=(
                    section.bottom_margin.inches if section.bottom_margin else None
                ),
                orientation=(str(section.orientation) if section.orientation else None),
            )
        )
    return sections


def _extract_header_footers(
    doc: DocumentObject,
) -> tuple[list[DocxHeaderFooter], list[DocxHeaderFooter]]:
    logger.debug("Extracting header/footer")
    headers = []
    footers = []
    for section in doc.sections:
        # Default
        if section.header and section.header.paragraphs:
            text = "\n".join(p.text for p in section.header.paragraphs)
            if text.strip():
                headers.append(DocxHeaderFooter(type="default", text=text))
        if section.footer and section.footer.paragraphs:
            text = "\n".join(p.text for p in section.footer.paragraphs)
            if text.strip():
                footers.append(DocxHeaderFooter(type="default", text=text))

        # First page
        if section.first_page_header and section.first_page_header.paragraphs:
            text = "\n".join(p.text for p in section.first_page_header.paragraphs)
            if text.strip():
                headers.append(DocxHeaderFooter(type="first_page", text=text))
        if section.first_page_footer and section.first_page_footer.paragraphs:
            text = "\n".join(p.text for p in section.first_page_footer.paragraphs)
            if text.strip():
                footers.append(DocxHeaderFooter(type="first_page", text=text))

        # Even page
        if section.even_page_header and section.even_page_header.paragraphs:
            text = "\n".join(p.text for p in section.even_page_header.paragraphs)
            if text.strip():
                headers.append(DocxHeaderFooter(type="even_page", text=text))
        if section.even_page_footer and section.even_page_footer.paragraphs:
            text = "\n".join(p.text for p in section.even_page_footer.paragraphs)
            if text.strip():
                footers.append(DocxHeaderFooter(type="even_page", text=text))
    return headers, footers


def _extract_paragraphs(doc: DocumentObject) -> list[DocxParagraph]:
    logger.debug("Extracting paragraphs")
    paragraphs = []
    for para in doc.paragraphs:
        runs = []
        for run in para.runs:
            runs.append(
                DocxRun(
                    text=run.text,
                    bold=run.bold,
                    italic=run.italic,
                    underline=run.underline,
                    font_name=run.font.name,
                    font_size=run.font.size.pt if run.font.size else None,
                    font_color=(
                        str(run.font.color.rgb)
                        if run.font.color and run.font.color.rgb
                        else None
                    ),
                )
            )
        paragraphs.append(
            DocxParagraph(
                text=para.text,
                style=para.style.name if para.style else None,
                alignment=str(para.alignment) if para.alignment else None,
                runs=runs,
            )
        )
    return paragraphs


def _extract_tables(doc: DocumentObject):
    logger.debug("Extracting tables")
    tables = []
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                cell_text = "\n".join(p.text for p in cell.paragraphs)
                row_data.append(cell_text)
            table_data.append(row_data)
        tables.append(table_data)
    return tables


def _extract_endnotes(file_like: io.BytesIO) -> list[DocxNote]:
    """Extracts endnotes which are like footnotes just that they are
    either forced to appear at the end of a section or at the end of a document"""
    logger.debug("Extracting endnotes")
    file_like.seek(0)
    endnotes = []
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    w_ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

    with zipfile.ZipFile(file_like, "r") as z:
        if "word/endnotes.xml" not in z.namelist():
            return endnotes

        with z.open("word/endnotes.xml") as f:
            tree = ET.parse(f)

        for en in tree.findall(".//w:endnote", ns):
            en_id = en.get(f"{w_ns}id") or ""
            if en_id not in ["-1", "0"]:  # Skip separator and continuation endnotes
                endnotes.append(
                    DocxNote(
                        id=en_id,
                        text="".join(t.text or "" for t in en.findall(".//w:t", ns)),
                    )
                )

    return endnotes


def _extract_formulas(file_like: io.BytesIO) -> list[DocxFormula]:
    """Extract all formulas from the document as LaTeX representations."""
    logger.debug("Extracting formulas")
    file_like.seek(0)
    doc = Document(file_like)
    formulas = []

    m_ns = "{http://schemas.openxmlformats.org/officeDocument/2006/math}"

    for element in doc.element.body.iter():
        tag = element.tag.split("}")[-1]

        # Display equation (oMathPara)
        if tag == "oMathPara":
            omath = element.find(f"{m_ns}oMath")
            if omath is not None:
                latex = _DocxFullTextExtractor.omml_to_latex(omath)
                if latex.strip():
                    formulas.append(DocxFormula(latex=latex, is_display=True))

        # Inline equation (oMath not inside oMathPara)
        elif tag == "oMath":
            # Check if parent is oMathPara - if so, skip (already handled above)
            parent = element.getparent()
            if parent is not None and parent.tag.split("}")[-1] == "oMathPara":
                continue
            latex = _DocxFullTextExtractor.omml_to_latex(element)
            if latex.strip():
                formulas.append(DocxFormula(latex=latex, is_display=False))

    return formulas


def read_docx(
    file_like: io.BytesIO, path: str | None = None
) -> Generator[DocxContent, Any, None]:
    """
    Extract all relevant content from a DOCX file.

    Args:
        file_like: A BytesIO object containing the DOCX file data.
        path: Optional file path to populate file metadata fields.

    Yields:
        MicrosoftDocxContent dataclass with all extracted content.
    """
    file_like.seek(0)
    doc = Document(file_like)

    # === Core Properties (Metadata) ===
    props = doc.core_properties
    metadata = DocxMetadata(
        title=props.title or "",
        author=props.author or "",
        subject=props.subject or "",
        keywords=props.keywords or "",
        category=props.category or "",
        comments=props.comments or "",
        created=(
            props.created.isoformat()
            if isinstance(props.created, datetime.datetime)
            else ""
        ),
        modified=(
            props.modified.isoformat()
            if isinstance(props.modified, datetime.datetime)
            else ""
        ),
        last_modified_by=props.last_modified_by or "",
        revision=props.revision,
    )

    # === Paragraphs ===
    paragraphs = _extract_paragraphs(doc=doc)

    # === Tables ===
    tables = _extract_tables(doc=doc)

    # === Headers and Footers ===
    headers, footers = _extract_header_footers(doc=doc)

    # === Images ===
    images = []
    for rel_id, rel in doc.part.rels.items():
        if "image" in rel.reltype:
            try:
                image_part = rel.target_part
                images.append(
                    DocxImage(
                        rel_id=rel_id,
                        filename=image_part.partname.split("/")[-1],
                        content_type=image_part.content_type,
                        data=io.BytesIO(image_part.blob),
                        size_bytes=len(image_part.blob),
                    )
                )
            except Exception as e:
                logger.debug(f"Image extraction failed for rel_id {rel_id} - {e}")
                images.append(DocxImage(rel_id=rel_id, error=str(e)))

    # === Hyperlinks ===
    hyperlinks = []
    rels = doc.part.rels
    for para in doc.paragraphs:
        for hyperlink in para._element.findall(
            ".//w:hyperlink",
            {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"},
        ):
            r_id = hyperlink.get(qn("r:id"))
            if r_id and r_id in rels and "hyperlink" in rels[r_id].reltype:
                text = "".join(
                    t.text or ""
                    for t in hyperlink.findall(
                        ".//w:t",
                        {
                            "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
                        },
                    )
                )
                hyperlinks.append(DocxHyperlink(text=text, url=rels[r_id].target_ref))

    # === Footnotes ===
    footnotes = _extract_footnotes(file_like=file_like)

    # === Endnotes ===
    endnotes = _extract_endnotes(file_like=file_like)

    # === Formulas ===
    formulas = _extract_formulas(file_like=file_like)

    # === Comments ===
    comments = _extract_comments(file_like=file_like)

    # === Sections (page layout) ===
    sections = _extract_sections(doc=doc)

    # === Styles used ===
    styles_set = set()
    for para in doc.paragraphs:
        if para.style:
            styles_set.add(para.style.name)
    styles = list(styles_set)

    # === Full text (convenience) ===
    full_text = _DocxFullTextExtractor.extract_full_text(
        file_like=file_like, include_formulas=True
    )
    base_full_text = _DocxFullTextExtractor.extract_full_text(
        file_like=file_like, include_formulas=False
    )

    metadata.populate_from_path(path)

    yield DocxContent(
        metadata=metadata,
        paragraphs=paragraphs,
        tables=tables,
        headers=headers,
        footers=footers,
        images=images,
        hyperlinks=hyperlinks,
        footnotes=footnotes,
        endnotes=endnotes,
        comments=comments,
        sections=sections,
        styles=styles,
        formulas=formulas,
        full_text=full_text,
        base_full_text=base_full_text,
    )
